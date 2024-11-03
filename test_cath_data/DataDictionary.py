"""
1、此模块为通过h5文件生成数据字典
2、其中数据来源只可以为：WIND,财汇，朝阳永续
3、生成的数据字典格式参考以上三个数据源的官方格式
4、生成的文件类型为docx文件
"""

import os
import traceback

import numpy as np
import pandas as pd

from docx import Document
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml



class Caihui_dictionary:
    pass


class Wind_dictionary:
    """
    Wind数据字典构建类
    """
    def __init__(self: object, file_name: str):
        """
        :param file_name: 文件名
        :type file_name : str
        """
        self.file_name = file_name

    def construct(self: object, hierarchy: int , document: object) -> None:
        """
        构建Wind数据字典
        :param hierarchy: 递归层级数
        :type hierarchy: int
        :param document: 文件对象
        :type document : object
        :return:
        """
        #一： 获取数据，表信息，字段信息
        print(f"开始生成{self.file_name}的数据字典")



class Zhaoyang_dictionary:
    """
    朝阳永续数据字典构建类
    """
    def __init__(self: object, file_name: str):
        """
        :param file_name: 文件名
        :type file_name: str
        """
        self.file_name = file_name

        # 定义一个表信息的队列，如果表信息显示改变，可从该处修改
        self.table_title_list = [("中文名称", "table_name_cn"), ("英文名称", "table_name"), ("表类型", "table_type"), ("主键", "key_name"), ("状态", "table_status"),
                                 ("唯一索引", "unique_index"), ("说明", "description")]

        # 定义一个字段信息的序列，如果字段信息显示改变，可从该处修改
        self.field_title_list = [("序号", ""), ("字段名", "field_name"), ("中文名", "field_name_cn"), ("类型", "field_type"), ("长度", "field_length"),
                                 ("注释", ""), ("说明", "")]

    def read_file(self: object, table_name: str)-> pd.DataFrame:
        """
        获取h5文件的内容
        :param table_name: 文件的sheet页
        :return:
        """
        store_client =  pd.HDFStore(self.file_name)
        df = store_client.get(table_name)
        store_client.close()
        return df

    def construct(self: object, hierarchy: int, document: object)-> None:
        """
        构建朝阳永续数据字典
        :param hierarchy: 递归层级数
        :type hierarchy: int
        :param document: 文件对象
        :type document : object
        :return:
        """
        # 一：获取数据；表信息，字段信息，说明信息
        print(f"开始生成{self.file_name}的数据字典")
        #1, 获取表信息
        res_table = self.read_file("describe_df")

        #2、再获取表字段信息
        res_field = self.read_file('table_field_df')
        res_field = res_field.drop_duplicates()
        res_field = res_field.replace(np.nan, "")

        #3、获取说明信息
        try:
            res_explan = self.read_file('fields_describe_df')
        except:
            print(traceback.format_exc())
            res_explan = pd.DataFrame()

        #二：将数据写入到docx文件中
        #1、写入标题
        title = document.add_heading("", level=4).add_run(res_table['table_name_cn'].values[0] + "[" + res_table['table_name'].values[0] + "]")
        title.font.name = u"宋体"
        title._element.rPr.rFonts.set(qn("w:eastAisa"), u"宋体")

        #2、写入表的基本信息
        table = document.add_table(rows=8, cols=2, style="Table Grid")
        row = table.rows[0]
        for cell in row.cells:
            shading_elm = parse_xml(r'<w:shd {} w:fill="#B0C4DE"/>'.format(nsdecls('w')))
            cell._tc.get_or_add_tcPr().append(shading_elm)
        first_start_lattice = table.cell(0, 0)
        first_start_lattice.text = "表基础信息"
        first_end_lattice = table.cell(0, 1)
        first_end_lattice.merge(first_start_lattice)
        for i in range(1, len(self.table_title_list)+1):
            table.cell(i, 0).text = self.table_title_list[i-1][0]
            table.cell(i, 1).text = res_table.loc[0,self.table_title_list[i-1][1]]

        #3、写入表字段的信息
        document.add_paragraph("").add_run("字段信息").bold = True
        table = document.add_table(rows=1, cols=len(self.field_title_list), style="Table Grid")
        row = table.rows[0]
        for cell in row.cells:
            shading_elm = parse_xml(r'<w:shd {} w:fill="#B0C4DE"/>'.format(nsdecls('w')))
            cell._tc.get_or_add_tcPr().append(shading_elm)
        for i in range(len(self.field_title_list)):
            table.cell(0, i).text = self.field_title_list[i][0]

        explan_list = [] if res_explan.empty else res_explan['field_name'].values.tolist()
        explan_s = []
        explan_count = 1
        for i in range(len(res_field)):
            row_cells = table.add_row().cells
            for x in range(len(self.field_title_list)):
                if self.field_title_list[x][0] == "序号":
                    row_cells[x].text = str(i + 1)
                elif self.field_title_list[x][0] == "注释":
                    row_cells[x].text = ""
                elif self.field_title_list[x][0] == "说明":
                    if res_field.loc[i, "field_name"] in explan_list:
                        row_cells[x].text = "说明" + str(explan_count)
                        explan_s.append(res_field.loc[i, "field_name"])
                        explan_s.append("说明" + str(explan_count))
                        explan_count += 1
                    else:
                        row_cells[x].text = ""
                else:
                    row_cells[x].text = res_field.loc[i, self.field_title_list[x][1]]

        #4、写入说明的信息
        if res_explan.empty:
            pass
        else:
            document.add_paragraph("").add_run("说明信息").bold = True
            for i in range(len(res_explan)):
                table = document.add_table(rows=4, cols=6, style="Table Grid")
                row = table.rows[0]
                for cell in row.cells:
                    cell._tc.get_or_add_tcPr().append(parse_xml(r'<w:shd {} w:fill="#B0C4DE"/>'.format(nsdecls('w'))))
                explan_start_lattice = table.cell(0, 0)
                explan_info = explan_s[explan_s.index(res_explan.loc[i, "field_name"]) + 1]
                explan_start_lattice.text = "字段详细内容" + "[" + explan_info + "]"
                explan_end_lattice = table.cell(0, 5)
                explan_end_lattice.merge(explan_start_lattice)

                col_name = res_explan.loc[i, "field_name"]
                col_name_cn = res_explan.loc[i, "field_name_cn"]
                key_value = res_explan.loc[i, "is_key"]
                if key_value == "0":
                    key_value = "否"
                else:
                    key_value = "是"
                field_type = res_explan.loc[i, "field_type"]
                table_name = res_explan.loc[i, "table_name"]
                field_description = res_explan.loc[i, "fileld_description"]

                # 下面这这部分写死了,感觉不够抽象
                table.cell(1, 0).text = "字段名称"
                table.cell(1, 1).text = col_name_cn
                table.cell(1, 2).text = "主键"
                table.cell(1, 3).text = key_value
                table.cell(1, 4).text = "类型"
                table.cell(1, 5).text = field_type

                table.cell(2, 0).text = "英文名称"
                table.cell(2, 1).text = col_name
                table.cell(2, 2).text = "所属对象"
                two_start_lattice = table.cell(2, 3)
                two_start_lattice.text = table_name
                two_end_lattice = table.cell(2, 5)
                two_end_lattice.merge(two_start_lattice)

                table.cell(3, 0).text = "释义"
                three_start_lattice = table.cell(3, 1)
                three_start_lattice.text = field_description
                three_end_lattice = table.cell(3, 5)
                three_end_lattice.merge(three_start_lattice)


def dirln_recursion(file_path: str, document: object, data_source: str, hierarchy: int = 1) -> None:
    """
    递归遍历当前目录下的所有目录和文件,并调用对应的方法
    :param file_path: 目录
    :type file_path: str 如: r"E:\朝阳永续数据\zhaoyang_datafile"
    :param document: 文档对象
    :type document: object
    :param data_source: 数据源
    :type data_source: str 如："zyyx"
    :param hierarchy : 递归层级数
    :type hierarchy: int 默认为1
    :return: None
    """
    file_list = os.listdir(file_path)
    for file_name in file_list:
        fileabpath = os.path.join(file_path, file_name)
        if os.path.isdir(fileabpath):
            if file_name == "__pycache__":
                pass
            else:
                hierarchy += 1
                title  = document.add_heading("", level= hierarchy-1).add_run(file_name)
                title.font.name = u"宋体"
                title._element.rPr.rFonts.set(qn('w:eastAsia'), u"宋体")
                dirln_recursion(fileabpath, document, data_source, hierarchy)
                hierarchy -= 1
        else:
            if file_name.endswith(".h5"):

                if data_source == "wind":
                    w_dict = Wind_dictionary(fileabpath)
                    w_dict.construct(hierarchy, document)
                elif data_source == "ch":
                    Caihui_dictionary()
                else:
                    zy_dict = Zhaoyang_dictionary(fileabpath)
                    zy_dict.construct(hierarchy, document)


def dictionary(file_path: str, file_name: str, data_source: str)-> dict:
    """
    1、构建数据字典的主函数
    :param file_path: 当前目录
    :type file_path: str  如：r"E:\朝阳永续数据\zhaoyang_datafile"
    :param file_name: 生成的数据字典文件名
    :type file_name : str 如："朝阳永续数据字典word版本.docx"
    :param data_source: 数据源
    :type data_source: str 如： "zyyx"
    :return: 执行结果
    :rtype return: dict
    """
    #try:
    document = Document(file_name)
    document.styles['Normal'].font.name = u"宋体"
    document.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
    dirln_recursion(file_path, document, data_source)
    document.save(file_name)
    result = {"code": 200, "message":"数据字典构建完毕"}
    # except:
    #     result = {"code": 500, "message":traceback.format_exc()}
    return result


if __name__ == "__main__":
    # 当前目录
    # 确保当前目录下的目录完整，h5文件也完整
    file_path = r"E:\万得数据\wind_datafile"
    # 生成的数据字典文件名称
    file_name = "wind数据字典汇总版.docx"
    # 数据源: 只可为wind,ch,zyyx
    data_source = "wind"
    res = dictionary(file_path, file_name, data_source)
    print(res)





