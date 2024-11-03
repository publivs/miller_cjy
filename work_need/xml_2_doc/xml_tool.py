import xml.etree.ElementTree as ET
from docx import Document

target_path = r'work_need\xml_2_doc\23北京银行012023-06-26.doc'

tree = ET.parse(target_path)
root = tree.getroot()
# 创建一个空的docx文档
doc = Document()

def process_node(node, parent):
    # 处理节点的标签
    tag = node.tag
    text = node.text
    if tag == "paragraph":
        # 创建段落
        p = doc.add_paragraph()
        if text:
            # 添加文本内容
            p.add_run(text)
    elif tag == "table":
        # 创建表格
        table = doc.add_table(rows=1, cols=1)
        if text:
            # 添加表格内容
            table.cell(0, 0).text = text
    # 处理节点的子节点
    for child in node:
        process_node(child, parent)

# 保存docx文件
doc.save('output.docx')