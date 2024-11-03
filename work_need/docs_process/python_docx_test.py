from docx import Document
from docx.shared import Inches
import docx

import pandas as pd

document = Document()
document.add_heading('Document Title', 0)

p = document.add_paragraph('A plain paragraph having some ')
p.add_run('bold').bold = True
p.add_run(' and some ')
p.add_run('italic.').italic = True

def df_to_word(df: pd.DataFrame,
                report_name:str) -> docx.Document:

    assert type(df) == pd.DataFrame, 'data has to be df'
    assert '.docx' in report_name, 'report_name has to be a .docx file'

    doc = docx.Document(report_name)

    table = doc.add_table(df.shape[0]+1, df.shape[1])

    for j in range(df.shape[-1]):
        table.cell(0,j).text = df.columns[j]

    for i in range(df.shape[0]):
        for j in range(df.shape[-1]):
            table.cell(i+1,j).text = str(df.values[i,j])

    doc.save(f'./{report_name}')

data = {
    "calorierbes": [420, 380, 390],
    "duratierbn": [50, 40, 45],
    "durationverg": [50, 40, 45],
    "duratiorgern": [50, 40, 45],
    "calorieers": [420, 380, 390],
    "calorierbers": [420, 380, 390],
    "calorierbes": [420, 380, 390]
        }

A = pd.DataFrame(data)
df_to_word(A,'test_report.docx')

{
    'product_name':'天玑航信宝集合资金信托计划',
    '月报日期':'202005',
    '项目名称':'天玑航信宝集合资金信托计划',



}
